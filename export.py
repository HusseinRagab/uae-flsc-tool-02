"""UAE FLSC 2018 - Word (.docx) and PDF exporters - multi-chapter."""
from __future__ import annotations

from io import BytesIO
from typing import List

from flsc_schema import ComplianceReport, Requirement, OCCUPANCY_DEFS, SectionBlock, DISCLAIMER
from figures import figure_caption, figures_for


def report_to_docx_bytes(r: ComplianceReport, kind: str = "detailed") -> bytes:
    if kind == "compact":
        return _compact_docx(r)
    return _detailed_docx(r)


def _add_docx_figure(doc, fig, width_in: float) -> None:
    from docx.shared import Inches, Pt
    path = fig["path"]
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    run = cap.add_run(figure_caption(fig))
    run.italic = True
    run.font.size = Pt(8)


def _compact_docx(r: ComplianceReport) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_heading("UAE FLSC 2018 - Compact report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("CDGH-OP-25, September 2018  |  Required system headers only").italic = True

    b = r.building
    if b.project_name:
        p = doc.add_paragraph()
        run = p.add_run(f"Project: {b.project_name}")
        run.bold = True

    doc.add_heading("Building Profile", level=1)
    profile = [
        ("Occupancy", b.occupancy),
        ("Height", f"{b.height_m} m  ({b.height_class})"),
        ("Storeys", f"{b.floors_above_grade} above + {b.floors_below_grade} basement"),
        ("GFA", f"{b.gross_floor_area_m2} m2"),
        ("Hazard class", b.hazard_class),
    ]
    tbl = doc.add_table(rows=len(profile), cols=2)
    tbl.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(profile):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = str(v)

    doc.add_heading("Required systems - headers only", level=1)
    for ch in r.chapters:
        items = [it for blk in ch.blocks for it in blk.items if it.status == "required"]
        if not items:
            continue
        doc.add_heading(f"{ch.chapter_code}  {ch.chapter_title}", level=2)
        for it in items:
            doc.add_paragraph(it.system, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Compact export. Use Detailed Word/PDF for figures and full clauses.").italic = True
    disc = doc.add_paragraph()
    disc.add_run(DISCLAIMER).italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _detailed_docx(r: ComplianceReport) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _fixed_table_layout(table, widths_cm):
        """Force Word to honour explicit column widths instead of auto-fitting
        to content. Two things are required: (1) w:tblLayout type='fixed', and
        (2) the w:tblGrid column widths — fixed layout renders from the GRID,
        not from per-cell widths, so without this the table falls back to
        equal columns."""
        tblPr = table._tbl.tblPr
        layout = tblPr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tblPr.append(layout)
        layout.set(qn("w:type"), "fixed")
        # 1 cm = 1440/2.54 twips. Set each grid column's width.
        grid = table._tbl.tblGrid
        cols = grid.findall(qn("w:gridCol"))
        for col, wcm in zip(cols, widths_cm):
            col.set(qn("w:w"), str(int(round(wcm * 1440 / 2.54))))

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)

    COL_SYSTEM = Cm(5.5)
    COL_SPEC = Cm(7.5)
    COL_CITE = Cm(4.4)

    title = doc.add_heading("UAE FIRE & LIFE SAFETY CODE 2018", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("CDGH-OP-25, September 2018  |  Detailed compliance report").italic = True

    b = r.building
    if b.project_name:
        p = doc.add_paragraph()
        run = p.add_run(f"Project: {b.project_name}")
        run.bold = True
        run.font.size = Pt(12)

    doc.add_heading("Building Profile", level=1)
    profile = [
        ("Occupancy", b.occupancy),
        ("Definition", OCCUPANCY_DEFS.get(b.occupancy, "")),
        ("Height", f"{b.height_m} m  ({b.height_class})"),
        ("Storeys", f"{b.floors_above_grade} above + {b.floors_below_grade} basement"),
        ("Plot area", f"{b.plot_area_m2} m2"),
        ("Ground-floor BUA", f"{b.ground_floor_bua_m2} m2"),
        ("Basement BUA", f"{b.basement_bua_m2} m2"),
        ("Total GFA", f"{b.gross_floor_area_m2} m2"),
        ("Hazard class", f"{b.hazard_class} (auto-derived)"),
    ]
    if b.has_high_ceiling:
        profile.append(("High ceiling", f"{b.ceiling_height_m} m"))
    if r.requires_wet_riser:
        profile.append(("Wet riser standpipes", str(b.wet_riser_standpipes)))
    tbl = doc.add_table(rows=len(profile), cols=2)
    tbl.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(profile):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = str(v)

    doc.add_heading("Chapter index - headers only. Full requirements follow.", level=1)
    idx = doc.add_table(rows=1, cols=3)
    idx.style = "Light Grid Accent 1"
    idx.rows[0].cells[0].text = "Section"
    idx.rows[0].cells[1].text = "Figure"
    idx.rows[0].cells[2].text = "Req / Rec / Cond / N/A"
    for cell in idx.rows[0].cells:
        for par in cell.paragraphs:
            for run in par.runs:
                run.bold = True
    for ch in r.chapters:
        req = rec = cond = na = 0
        for blk in ch.blocks:
            for it in blk.items:
                if it.status == "required":
                    req += 1
                elif it.status == "recommended":
                    rec += 1
                elif it.status == "conditional":
                    cond += 1
                else:
                    na += 1
        figs = figures_for(ch.chapter_code)
        fig_bits = []
        for f in figs:
            bit = f["figure"]
            if f["page"]:
                bit += f"  p.{f['page']}"
            fig_bits.append(bit)
        row = idx.add_row().cells
        row[0].text = f"{ch.chapter_code}  {ch.chapter_title}"
        row[1].text = "  |  ".join(fig_bits) or "-"
        row[2].text = f"{req} / {rec} / {cond} / {na}"
    note = doc.add_paragraph()
    note.add_run("Each following section is one system: authentic UAE FLSC 2018 figure, then the evaluated requirements.").italic = True

    def req_block(title: str, items: List[Requirement]):
        if not items:
            return
        doc.add_heading(title, level=1)
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        t.autofit = False
        t.allow_autofit = False
        _fixed_table_layout(t, [5.5, 7.5, 4.4])
        hdr = t.rows[0].cells
        hdr[0].text = "System"; hdr[1].text = "Spec / Detail"; hdr[2].text = "Code Ref"
        for cell in hdr:
            for par in cell.paragraphs:
                for run in par.runs:
                    run.bold = True
        for req in items:
            row = t.add_row().cells
            tag = "" if req.status == "required" else f" [{req.status}]"
            row[0].text = f"{req.system}{tag}"
            parts = []
            if req.spec:    parts.append(f"Spec: {req.spec}")
            if req.detail:  parts.append(req.detail)
            row[1].text = "\n\n".join(parts) if parts else "-"
            cite = " - ".join(p for p in (req.code_ref, req.page_ref) if p)
            row[2].text = cite or "-"
        # Word column widths must be set on every cell of the column to stick.
        for row in t.rows:
            row.cells[0].width = COL_SYSTEM
            row.cells[1].width = COL_SPEC
            row.cells[2].width = COL_CITE

    for ch in r.chapters:
        doc.add_heading(f"{ch.chapter_code} - {ch.chapter_title}", level=1)
        if ch.selected_branch:
            p = doc.add_paragraph()
            p.add_run("Matched branch: ").bold = True
            p.add_run(f"{ch.selected_branch}  -  {ch.selected_branch_section}")
        for fig in figures_for(ch.chapter_code):
            _add_docx_figure(doc, fig, width_in=6.4)
        for block in ch.blocks:
            req_block(block.title, block.items)

    if r.high_ceiling and r.high_ceiling.applies:
        hc = r.high_ceiling
        doc.add_heading("FP - High Ceiling Sprinkler Design (Table 9.29.A)", level=1)
        rows = [
            ("Ceiling height", f"{hc.ceiling_height_m} m"),
            ("Hazard class", hc.hazard_class),
            ("Height band", hc.height_range or "out of tabulated range"),
            ("K-factor", hc.k_factor or "-"),
            ("Min pressure", hc.min_pressure or "-"),
            ("Min sprinklers", str(hc.min_sprinklers) if hc.min_sprinklers else "-"),
            ("Density", hc.density or "-"),
            ("Design area", hc.design_area or "-"),
            ("Pump (no hydrant)", f"{hc.pump_without_hydrant_gpm} gpm" if hc.pump_without_hydrant_gpm else "-"),
            ("Pump (with hydrant)", f"{hc.pump_with_hydrant_gpm} gpm" if hc.pump_with_hydrant_gpm else "-"),
        ]
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = "Light Grid Accent 1"
        for i, (k, v) in enumerate(rows):
            t.rows[i].cells[0].text = k
            t.rows[i].cells[1].text = v
        if hc.note:
            doc.add_paragraph().add_run(f"Note: {hc.note}").italic = True

    doc.add_paragraph()
    doc.add_heading("Disclaimer", level=2)
    disc_p = doc.add_paragraph()
    disc_run = disc_p.add_run(DISCLAIMER)
    disc_run.italic = True
    disc_run.font.size = Pt(9)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = foot.add_run("Generated by UAE FLSC compliance tool. Not an official Civil Defence document.")
    run.italic = True
    run.font.size = Pt(8)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def report_to_pdf_bytes(r: ComplianceReport, kind: str = "detailed") -> bytes:
    if kind == "compact":
        return _compact_pdf(r)
    return _detailed_pdf(r)


def _rl_figure(fig, max_w, max_h):
    from reportlab.platypus import Image as RLImage, Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib import colors

    path = fig["path"]
    if not path.exists():
        return []
    img = RLImage(str(path))
    iw, ih = float(img.imageWidth), float(img.imageHeight)
    scale = min(max_w / iw, max_h / ih)
    img.drawWidth = iw * scale
    img.drawHeight = ih * scale
    cap = ParagraphStyle(
        "figcap",
        fontName="Helvetica-Oblique",
        fontSize=8,
        textColor=colors.HexColor("#163A3A"),
        leading=10,
        spaceAfter=8,
    )
    return [img, Spacer(1, 4), Paragraph(figure_caption(fig).replace("—", "-"), cap)]


def _compact_pdf(r: ComplianceReport) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, ListFlowable, ListItem,
    )

    TEAL = colors.HexColor("#163A3A")
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.7 * cm, rightMargin=1.7 * cm,
        topMargin=1.6 * cm, bottomMargin=1.5 * cm,
        title="UAE FLSC compact report",
    )
    ss = getSampleStyleSheet()
    title_s = ParagraphStyle("t", parent=ss["Title"], fontSize=16, textColor=TEAL, spaceAfter=4)
    sub_s = ParagraphStyle("s", parent=ss["Normal"], fontSize=9, textColor=colors.grey, spaceAfter=8)
    h1 = ParagraphStyle("h1", parent=ss["Heading1"], fontSize=11, textColor=TEAL, spaceBefore=10, spaceAfter=4)
    normal = ParagraphStyle("n", parent=ss["Normal"], fontSize=9, leading=12)
    small = ParagraphStyle("small", parent=ss["Normal"], fontSize=8, textColor=colors.dimgrey, leading=10)

    b = r.building
    story = [
        Paragraph("UAE FIRE & LIFE SAFETY CODE 2018", title_s),
        Paragraph("CDGH-OP-25 | September 2018 | Compact report - required systems", sub_s),
        Paragraph(f"<b>{b.project_name or 'Fire & life safety report'}</b>", normal),
        Spacer(1, 8),
        Paragraph("Building profile", h1),
    ]
    profile = [
        ["Occupancy", b.occupancy],
        ["Height", f"{b.height_m} m ({b.height_class})"],
        ["Storeys", f"{b.floors_above_grade} above + {b.floors_below_grade} basement"],
        ["GFA", f"{b.gross_floor_area_m2} m2"],
        ["Hazard class", b.hazard_class],
    ]
    t = Table(profile, colWidths=[4.5 * cm, 12 * cm])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(t)

    req_n = rec_n = cond_n = na_n = 0
    for ch in r.chapters:
        for blk in ch.blocks:
            for it in blk.items:
                if it.status == "required":
                    req_n += 1
                elif it.status == "recommended":
                    rec_n += 1
                elif it.status == "conditional":
                    cond_n += 1
                else:
                    na_n += 1
    story.append(Spacer(1, 8))
    story.append(Paragraph("Requirement counts", h1))
    story.append(Paragraph(
        f"<b>Required {req_n}</b> &nbsp; Recommended {rec_n} &nbsp; "
        f"Conditional {cond_n} &nbsp; Not required {na_n}",
        normal,
    ))
    story.append(Spacer(1, 6))
    story.append(Paragraph("Required systems - headers only", h1))

    for ch in r.chapters:
        items = [it for blk in ch.blocks for it in blk.items if it.status == "required"]
        if not items:
            continue
        story.append(Paragraph(f"{ch.chapter_code} &nbsp; {ch.chapter_title}", h1))
        bullets = [ListItem(Paragraph(it.system, normal), leftIndent=12) for it in items]
        story.append(ListFlowable(bullets, bulletType="bullet", start="•"))

    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="100%", thickness=0.4, color=colors.grey))
    story.append(Paragraph(
        "Compact export. Use Detailed PDF for figures and full clauses.",
        small,
    ))
    story.append(Paragraph(f"<i>{DISCLAIMER}</i>", small))
    doc.build(story)
    return buf.getvalue()



def _detailed_pdf(r: ComplianceReport) -> bytes:
    """Portrait A4: cover + chapter index, then figure + requirements per chapter."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak,
    )
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT

    TEAL = colors.HexColor("#163A3A")
    PAPER = colors.HexColor("#F3EFE6")
    INK = colors.HexColor("#1A1F1C")
    MUTED = colors.HexColor("#5C635E")
    RULE = colors.HexColor("#D9D1C6")
    WHITE = colors.white
    STATUS_COL = {
        "required": colors.HexColor("#8B1E1E"),
        "recommended": colors.HexColor("#2C4A6E"),
        "conditional": colors.HexColor("#6B4E16"),
        "not_required": colors.HexColor("#5C635E"),
    }
    STATUS_LAB = {
        "required": "REQUIRED",
        "recommended": "RECOMMENDED",
        "conditional": "CONDITIONAL",
        "not_required": "NOT REQUIRED",
    }

    W, H = A4
    b = r.building
    occ_label = OCCUPANCY_DEFS.get(b.occupancy, b.occupancy).split(" - ")[0]
    hc = str(b.height_class).replace("_", " ")

    def _counts(ch):
        req = rec = cond = na = 0
        for blk in ch.blocks:
            for it in blk.items:
                if it.status == "required":
                    req += 1
                elif it.status == "recommended":
                    rec += 1
                elif it.status == "conditional":
                    cond += 1
                else:
                    na += 1
        return req, rec, cond, na

    tot = {"required": 0, "recommended": 0, "conditional": 0, "not_required": 0}
    for ch in r.chapters:
        for blk in ch.blocks:
            for it in blk.items:
                tot[it.status] = tot.get(it.status, 0) + 1

    def _later(c, doc):
        c.saveState()
        c.setFillColor(TEAL)
        c.rect(0, H - 8, W, 8, fill=1, stroke=0)
        c.setStrokeColor(RULE)
        c.setLineWidth(0.4)
        c.line(48, 28, W - 48, 28)
        c.setFillColor(MUTED)
        c.setFont("Helvetica", 7)
        name = (b.project_name or "Fire & life safety report")[:70]
        c.drawString(48, H - 22, name)
        c.drawString(48, 16, "UAE FLSC 2018  |  CDGH-OP-25  |  Design aid - not a Civil Defence submission")
        c.drawRightString(W - 48, 16, f"Page {doc.page}")
        c.restoreState()

    def _first(c, doc):
        c.saveState()
        c.setFillColor(PAPER)
        c.rect(0, 0, W, H, fill=1, stroke=0)
        c.setFillColor(TEAL)
        c.rect(0, H - 100, W, 100, fill=1, stroke=0)
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 15)
        c.drawString(48, H - 52, "UAE FIRE & LIFE SAFETY CODE 2018")
        c.setFillColor(colors.HexColor("#C7DBDB"))
        c.setFont("Helvetica", 9)
        c.drawString(48, H - 72, "CDGH-OP-25  |  September 2018  |  Detailed compliance report")
        c.restoreState()
        _later(c, doc)

    h_title = ParagraphStyle("ptitle", fontName="Helvetica-Bold", fontSize=16, textColor=INK, spaceAfter=10, leading=20)
    h_sec = ParagraphStyle("psec", fontName="Helvetica-Bold", fontSize=11, textColor=TEAL, spaceBefore=8, spaceAfter=6)
    h_ch = ParagraphStyle("pch", fontName="Helvetica-Bold", fontSize=11, textColor=WHITE, leading=14)
    k_s = ParagraphStyle("pk", fontName="Helvetica-Bold", fontSize=9, textColor=MUTED, leading=12)
    v_s = ParagraphStyle("pv", fontName="Helvetica", fontSize=10, textColor=INK, leading=13)
    n_s = ParagraphStyle("pn", fontName="Helvetica", fontSize=9, textColor=INK, leading=12)
    m_s = ParagraphStyle("pm", fontName="Helvetica", fontSize=8.5, textColor=MUTED, leading=11)
    i_s = ParagraphStyle("pi", fontName="Helvetica-Oblique", fontSize=8, textColor=MUTED, leading=10)
    sys_s = ParagraphStyle("psys", fontName="Helvetica-Bold", fontSize=10, textColor=INK, leading=13)
    st_s = ParagraphStyle("pst", fontName="Helvetica-Bold", fontSize=7, alignment=TA_RIGHT, leading=10)
    idx_h = ParagraphStyle("pidh", fontName="Helvetica-Bold", fontSize=7, textColor=WHITE, leading=9)
    idx_b = ParagraphStyle("pidb", fontName="Helvetica-Bold", fontSize=8.5, textColor=INK, leading=11)
    idx_m = ParagraphStyle("pidm", fontName="Helvetica", fontSize=7, textColor=MUTED, leading=9)

    story = [Spacer(1, 62)]
    story.append(Paragraph(b.project_name or "Fire & life safety report", h_title))
    story.append(Paragraph("Building profile", h_sec))

    rows = [
        ["Occupancy", occ_label],
        ["Height class", f"{b.height_m} m  |  {hc}"],
        ["Storeys", f"{b.floors_above_grade} above grade  |  {b.floors_below_grade} basement"],
        ["Gross floor area", f"{b.gross_floor_area_m2:,.0f} m2"],
        ["GF / basement BUA", f"{b.ground_floor_bua_m2:,.0f} m2  |  {b.basement_bua_m2:,.0f} m2"],
        ["Plot area", f"{b.plot_area_m2:,.0f} m2"],
        ["Hazard class", b.hazard_class],
    ]
    if b.has_high_ceiling:
        rows.append(["High ceiling", f"{b.ceiling_height_m} m"])
    if r.requires_wet_riser:
        rows.append(["Wet riser standpipes", str(b.wet_riser_standpipes)])
    prof = [[Paragraph(k, k_s), Paragraph(str(v), v_s)] for k, v in rows]
    pt = Table(prof, colWidths=[4.4 * cm, 12.4 * cm])
    pt.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
    ]))
    story.append(pt)
    story.append(Spacer(1, 8))
    story.append(HRFlowable(width="100%", thickness=0.5, color=RULE, spaceAfter=10))
    story.append(Paragraph("Requirement counts", h_sec))

    chips = []
    for st, label in (("required", "REQUIRED"), ("recommended", "RECOMMENDED"),
                      ("conditional", "CONDITIONAL"), ("not_required", "NOT REQUIRED")):
        inner = Table(
            [[Paragraph(f'<font color="{STATUS_COL[st].hexval()}"><b>{label}</b></font>', n_s)],
             [Paragraph(f"<b>{tot[st]}</b>", ParagraphStyle("cn", fontName="Helvetica-Bold", fontSize=14, textColor=INK))]],
            colWidths=[3.9 * cm],
        )
        inner.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 1, STATUS_COL[st]),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("BACKGROUND", (0, 0), (-1, -1), WHITE),
        ]))
        chips.append(inner)
    ct = Table([chips], colWidths=[4.2 * cm] * 4)
    ct.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 2),
        ("RIGHTPADDING", (0, 0), (-1, -1), 2),
    ]))
    story.append(ct)
    story.append(Spacer(1, 10))
    story.append(Paragraph("Chapter index  -  headers only. Full requirements follow.", h_sec))

    idx = [[Paragraph("Section", idx_h), Paragraph("Figure", idx_h), Paragraph("Req / Rec / Cond / N/A", idx_h)]]
    for ch in r.chapters:
        figs = figures_for(ch.chapter_code)
        fig_bits = []
        for f in figs:
            bit = f["figure"]
            if f["page"]:
                bit += f"  p.{f['page']}"
            fig_bits.append(bit)
        fig_line = "  |  ".join(fig_bits) or "-"
        req, rec, cond, na = _counts(ch)
        idx.append([
            Paragraph(f"{ch.chapter_code}  {ch.chapter_title}", idx_b),
            Paragraph(fig_line, idx_m),
            Paragraph(f"{req} / {rec} / {cond} / {na}", idx_b),
        ])
    it = Table(idx, colWidths=[7.2 * cm, 6.2 * cm, 3.4 * cm])
    it.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), TEAL),
        ("BACKGROUND", (0, 1), (-1, -1), WHITE),
        ("BOX", (0, 0), (-1, -1), 0.4, RULE),
        ("INNERGRID", (0, 1), (-1, -1), 0.3, RULE),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(it)
    story.append(Spacer(1, 8))
    story.append(Paragraph(
        "Each following page is one system: authentic UAE FLSC 2018 figure, then the evaluated requirements.",
        m_s,
    ))
    story.append(Spacer(1, 4))
    story.append(Paragraph(DISCLAIMER, i_s))

    story.append(PageBreak())

    inner_w = 16.8 * cm

    for ch in r.chapters:
        bar = Table([[Paragraph(f"{ch.chapter_code}  |  {ch.chapter_title}", h_ch)]], colWidths=[inner_w])
        bar.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), TEAL),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        story.append(bar)
        story.append(Spacer(1, 6))
        if ch.selected_branch:
            story.append(Paragraph(
                f"Matched branch: {ch.selected_branch} | {ch.selected_branch_section}",
                i_s,
            ))
            story.append(Spacer(1, 4))
        for fig in figures_for(ch.chapter_code):
            story.extend(_rl_figure(fig, inner_w, 8.2 * cm))
        for blk in ch.blocks:
            story.append(Paragraph(blk.title, h_sec))
            for req in blk.items:
                stcol = STATUS_COL.get(req.status, MUTED)
                lab = STATUS_LAB.get(req.status, req.status.upper())
                bits = [[Paragraph(f'<font color="{stcol.hexval()}">{lab}</font>', st_s)]]
                bits.append([Paragraph(req.system, sys_s)])
                if req.spec:
                    bits.append([Paragraph(req.spec.replace("\n", "<br/>"), n_s)])
                if req.detail:
                    bits.append([Paragraph(req.detail.replace("\n", "<br/>"), m_s)])
                cite = "  |  ".join(p for p in (req.code_ref, req.page_ref) if p)
                if cite:
                    bits.append([Paragraph(cite, i_s)])
                item = Table(bits, colWidths=[inner_w])
                item.setStyle(TableStyle([
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 1),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                    ("LINEBEFORE", (0, 0), (0, -1), 3, stcol),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]))
                story.append(item)
                story.append(Spacer(1, 6))
        story.append(Spacer(1, 8))

    if r.high_ceiling and r.high_ceiling.applies:
        hc_ = r.high_ceiling
        bar = Table([[Paragraph("FP  |  High-ceiling sprinkler design  -  Table 9.29.A", h_ch)]], colWidths=[inner_w])
        bar.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), TEAL),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        story.append(bar)
        story.append(Paragraph(
            f"{hc_.ceiling_height_m} m  |  {hc_.hazard_class}  |  {hc_.height_range or 'out of tabulated range'}",
            n_s,
        ))
        if hc_.note:
            story.append(Paragraph(hc_.note, i_s))

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=48, rightMargin=48,
        topMargin=52, bottomMargin=42,
        title=f"{b.project_name or 'UAE FLSC'} - detailed Fire & Life Safety Report",
    )
    doc.build(story, onFirstPage=_first, onLaterPages=_later)
    return buf.getvalue()
